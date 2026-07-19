"""
ui/app.py — Streamlit Web UI · Lex Modernus design
Run from project root: streamlit run ui/app.py
"""

import sys, os, tempfile, threading, io
import streamlit as st
import mistune
from docx import Document
from docx.shared import Pt, Inches

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
if ROOT not in sys.path:
    sys.path.insert(0, ROOT)

from config import (
    OUTPUT_DIGEST_PATH, INTERMEDIATE_DIR,
    DEFAULT_WINDOW, DEFAULT_MIN_GAP, VERBATIM_WINDOW,
    REDUCE_MAX_TOKENS,
)
from models.constants import BAR_SUBJECTS, KEYWORDS
from swarm.swarm import ProviderSwarm
from swarm.provider import Provider
from pipeline.p00_fetch import fetch_case_from_url, ingest_pdf
from pipeline.p02_scan import build_automaton, scan_lines, deduplicate_hits, pre_extract_verbatim
from pipeline.p03_chunk import build_all_payloads
from pipeline.p04_map import process_all_chunks
from pipeline.p05_stitch import build_compiled_stream
from pipeline.p06_reduce import run_reduce

# ── Page config — MUST be first Streamlit call ────────────────────────────────
st.set_page_config(
    page_title="Lex Modernus | Philippine Case Digest AI",
    page_icon="⚖️",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ── Full CSS override — forces light mode regardless of system/browser theme ──
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Source+Serif+4:ital,wght@0,400;0,600;0,700;1,400&family=Hanken+Grotesk:wght@400;500;600;700&family=JetBrains+Mono:wght@400&display=swap');

/* ═══ FORCE LIGHT MODE — override every Streamlit dark-mode selector ═══ */
:root,
html,
body,
[data-testid="stAppViewContainer"],
[data-testid="stAppViewBlockContainer"],
[data-testid="stVerticalBlock"],
[data-testid="stMain"],
.main,
.block-container,
section[data-testid="stSidebar"],
[data-testid="stSidebar"] > div,
[data-testid="stSidebar"] > div > div,
div[class*="st-"],
div[class*="css-"] {
    background-color: #f8fafc !important;
    color: #1e293b !important;
    font-family: 'Hanken Grotesk', sans-serif !important;
}

/* Main content area specifically */
[data-testid="stAppViewContainer"] {
    background: #f8fafc !important;
}
[data-testid="stMain"] > div {
    background: #f8fafc !important;
}
.block-container {
    background: #f8fafc !important;
    padding-top: 2rem !important;
    max-width: 1100px !important;
}

/* ═══ HIDE STREAMLIT CHROME ═══ */
#MainMenu { visibility: hidden !important; }
footer { visibility: hidden !important; }
[data-testid="stToolbar"] { display: none !important; }
[data-testid="stDecoration"] { display: none !important; }

/* ═══ INPUTS ═══ */
[data-testid="stTextInput"] input {
    background: #f8fafc !important;
    border: 2.5px solid #64748b !important;
    border-radius: 8px !important;
    color: #1e293b !important;
    font-family: 'Hanken Grotesk', sans-serif !important;
    font-size: 14px !important;
    padding: 12px 16px !important;
}
[data-testid="stTextInput"] input:focus {
    border-color: #D4AF37 !important;
    box-shadow: 0 0 0 3px rgba(212,175,55,0.15) !important;
    outline: none !important;
}
[data-testid="stTextInput"] input::placeholder {
    color: #64748b !important;
}
[data-testid="stTextInput"] label {
    color: #64748b !important;
    font-size: 11px !important;
    font-weight: 700 !important;
    letter-spacing: 0.08em !important;
    text-transform: uppercase !important;
}

/* ═══ SELECTBOX ═══ */
[data-testid="stSelectbox"] > div > div {
    background: #f8fafc !important;
    border: 2px solid #94a3b8 !important;
    border-radius: 8px !important;
    color: #1e293b !important;
}
[data-testid="stSelectbox"] label {
    color: #64748b !important;
    font-size: 11px !important;
    font-weight: 700 !important;
    letter-spacing: 0.08em !important;
    text-transform: uppercase !important;
}

/* ═══ RADIO ═══ */
[data-testid="stRadio"] {
    background: transparent !important;
}
[data-testid="stRadio"] > label {
    color: #64748b !important;
    font-size: 11px !important;
    font-weight: 700 !important;
    letter-spacing: 0.08em !important;
    text-transform: uppercase !important;
}
[data-testid="stRadio"] div[role="radio"] {
    border-color: #475569 !important;
    background: #ffffff !important;
    border-width: 2px !important;
}
[data-testid="stRadio"] div[role="radio"][aria-checked="true"] {
    border-color: #D4AF37 !important;
    background: #ffffff !important;
    box-shadow: 0 0 0 3px rgba(212,175,55,0.2) !important;
}
[data-testid="stRadio"] div[role="radio"][aria-checked="true"]::after {
    background: #D4AF37 !important;
}
[data-testid="stRadio"] label {
    color: #1e293b !important;
    font-size: 14px !important;
    font-weight: 500 !important;
}
/* Highlight the entire selected radio row */
[data-testid="stRadio"] label[data-testid="stRadioLabel"][aria-checked="true"] {
    background: rgba(212,175,55,0.08) !important;
}

/* ═══ FILE UPLOADER ═══ */
[data-testid="stFileUploader"] {
    background: #ffffff !important;
    border: 2px dashed #94a3b8 !important;
    border-radius: 8px !important;
}
[data-testid="stFileUploader"] label {
    color: #64748b !important;
}

/* ═══ PRIMARY BUTTON — Analyze Case ═══ */
div[data-testid="stButton"] > button[kind="primary"],
div[data-testid="stButton"] > button[kind="primary"]:focus {
    background: #001B3D !important;
    color: #ffffff !important;
    border: none !important;
    border-radius: 8px !important;
    font-family: 'Hanken Grotesk', sans-serif !important;
    font-weight: 700 !important;
    font-size: 14px !important;
    letter-spacing: 0.04em !important;
    padding: 14px 28px !important;
    width: 100% !important;
    box-shadow: 0 4px 14px rgba(0,27,61,0.3) !important;
    transition: all 0.2s !important;
}
div[data-testid="stButton"] > button[kind="primary"]:hover {
    background: #0f2d5a !important;
    transform: translateY(-1px) !important;
    box-shadow: 0 6px 20px rgba(0,27,61,0.35) !important;
}
/* Force all inner button elements to use dark background + white text */
div[data-testid="stButton"] button[kind="primary"] * {
    background: transparent !important;
    color: #ffffff !important;
}

/* ═══ DOWNLOAD BUTTON ═══ */
[data-testid="stDownloadButton"] button {
    background: transparent !important;
    border: 1.5px solid #D4AF37 !important;
    color: #001B3D !important;
    font-weight: 600 !important;
    border-radius: 8px !important;
    font-family: 'Hanken Grotesk', sans-serif !important;
    transition: all 0.2s !important;
}
[data-testid="stDownloadButton"] button:hover {
    background: rgba(212,175,55,0.1) !important;
}

/* ═══ ALERTS ═══ */
[data-testid="stAlert"] {
    border-radius: 8px !important;
}

/* ═══ DIVIDER ═══ */
hr {
    border-color: #e2e8f0 !important;
}

/* ═══ CUSTOM COMPONENTS ═══ */
.lex-hero {
    text-align: center;
    padding: 40px 0 32px;
}
.lex-hero h2 {
    font-family: 'Source Serif 4', serif !important;
    font-size: 42px !important;
    font-weight: 700 !important;
    color: #001B3D !important;
    margin: 0 0 10px 0 !important;
    letter-spacing: -0.01em !important;
    line-height: 1.15 !important;
}
.lex-hero p {
    font-size: 17px !important;
    color: #64748b !important;
    font-weight: 400 !important;
    max-width: 540px !important;
    margin: 0 auto 20px !important;
    line-height: 1.6 !important;
}
.flag-bar {
    display: flex;
    justify-content: center;
    gap: 8px;
    margin-top: 16px;
}
.flag-bar span {
    display: block;
    height: 4px;
    width: 52px;
    border-radius: 2px;
}

.lex-brand {
    padding: 28px 24px 18px;
    border-bottom: 1px solid #e2e8f0;
    margin-bottom: 24px;
    background: #ffffff !important;
}
.lex-brand h1 {
    font-family: 'Source Serif 4', serif !important;
    font-size: 22px !important;
    font-weight: 700 !important;
    color: #D4AF37 !important;
    margin: 8px 0 3px !important;
}
.lex-brand p {
    font-size: 9px !important;
    letter-spacing: 0.15em !important;
    text-transform: uppercase !important;
    color: #64748b !important;
    font-weight: 700 !important;
    margin: 0 !important;
}
.gold-rule {
    height: 2px;
    background: linear-gradient(90deg, #D4AF37, transparent);
    margin-top: 14px;
    border: none;
}

.glass-card {
    background: #ffffff !important;
    border: 1px solid #e2e8f0 !important;
    border-radius: 12px !important;
    padding: 28px !important;
    box-shadow: 0 2px 12px rgba(0,0,0,0.05) !important;
    margin-bottom: 24px !important;
}
.section-label {
    font-size: 11px;
    letter-spacing: 0.08em;
    text-transform: uppercase;
    color: #64748b;
    font-weight: 700;
    margin-bottom: 10px;
    display: flex;
    align-items: center;
    gap: 6px;
}

.bento-grid {
    display: grid;
    grid-template-columns: repeat(3, 1fr);
    gap: 18px;
    margin-top: 8px;
}
.bento-card {
    background: #ffffff !important;
    border: 1px solid #e2e8f0 !important;
    border-radius: 12px !important;
    padding: 24px !important;
    transition: border-color 0.2s, transform 0.2s;
}
.bento-card:hover {
    border-color: rgba(212,175,55,0.6) !important;
    transform: translateY(-3px);
}
.bento-icon {
    width: 44px; height: 44px;
    background: #f1f5f9;
    border-radius: 10px;
    display: flex; align-items: center; justify-content: center;
    margin-bottom: 14px;
    font-size: 22px;
}
.bento-card h3 {
    font-family: 'Source Serif 4', serif !important;
    font-size: 17px !important;
    font-weight: 600 !important;
    color: #001B3D !important;
    margin: 0 0 8px !important;
}
.bento-card p {
    font-size: 13px !important;
    color: #64748b !important;
    line-height: 1.6 !important;
    margin: 0 !important;
}

.stage-row { display: flex; flex-wrap: wrap; gap: 6px; margin-bottom: 10px; }
.stage-pill {
    display: inline-flex; align-items: center; gap: 4px;
    padding: 4px 11px; border-radius: 99px;
    font-size: 11px; font-weight: 600;
    font-family: 'Hanken Grotesk', sans-serif;
}
.s-done   { background:#e6f4ea; color:#276239; border:1px solid #b7dfc3; }
.s-active { background:#fff8e1; color:#7a5c00; border:1px solid #f0d580; }
.s-wait   { background:#f1f5f9; color:#94a3b8; border:1px solid #e2e8f0; }
.stage-msg { font-size:12px; color:#64748b; font-style:italic; }

.digest-output {
    font-family: 'Source Serif 4', serif !important;
    font-size: 15px !important;
    line-height: 1.8 !important;
    color: #1e293b !important;
    background: #ffffff !important;
    border: 1px solid #e2e8f0 !important;
    border-radius: 12px !important;
    padding: 32px 40px !important;
    box-shadow: 0 2px 12px rgba(0,0,0,0.04) !important;
}
.digest-output h2 {
    font-size: 16px !important; font-weight: 700 !important;
    color: #001B3D !important; margin-top: 24px !important;
    border-bottom: 1px solid #e2e8f0 !important; padding-bottom: 6px !important;
}

.lex-footer {
    margin-top: 60px;
    padding-top: 18px;
    border-top: 1px solid #e2e8f0;
    display: flex;
    justify-content: space-between;
    align-items: center;
    font-size: 11px;
    color: #94a3b8;
    font-family: 'JetBrains Mono', monospace;
}
</style>

""", unsafe_allow_html=True)

# ════════════════════════════════════════════════════════════════════
# MAIN
# ════════════════════════════════════════════════════════════════════

# Hero
st.markdown("""
<div class="lex-hero">
    <h2>Philippine Case Digest AI</h2>
    <p>Advanced Jurisprudence Analysis Engine for the Philippine Bar and Judiciary.</p>
    <div class="flag-bar">
        <span style="background:#0038a8;"></span>
        <span style="background:#ce1126;"></span>
        <span style="background:#D4AF37;"></span>
    </div>
</div>
""", unsafe_allow_html=True)

# Input card
st.markdown('<div class="glass-card">', unsafe_allow_html=True)
st.markdown('<div class="section-label">🔗 &nbsp; Case Source</div>', unsafe_allow_html=True)

input_type = st.radio("Input Type", ["URL", "PDF Upload"], index=0, horizontal=True)

if input_type == "URL":
    url_input = st.text_input(
        "url", label_visibility="collapsed",
        placeholder="https://elibrary.judiciary.gov.ph/thebookshelf/showdocs/...",
    )
    pdf_file = None
else:
    url_input = ""
    pdf_file = st.file_uploader("pdf", type=["pdf"], label_visibility="collapsed")

bar_subject = st.selectbox("Bar Subject", BAR_SUBJECTS, index=0)

generate_clicked = st.button("⚡ Analyze Case", type="primary", use_container_width=True)

st.markdown("""
<p style='font-size:10px;color:#94a3b8;text-align:center;font-style:italic;margin-top:2px;
          font-family:"JetBrains Mono",monospace;'>
    Swarm built on click · No idle API calls
</p>
""", unsafe_allow_html=True)
st.markdown('</div>', unsafe_allow_html=True)

# Progress + output placeholders
progress_area = st.empty()
output_area   = st.empty()

# Bento cards (shown until first result)
if not st.session_state.get("digest"):
    st.markdown("""
    <div class="bento-grid">
        <div class="bento-card">
            <div class="bento-icon">📚</div>
            <h3>Jurisprudence</h3>
            <p>Access Supreme Court decisions with cross-referenced citations and relevant doctrines from the SC E-Library and Lawphil.</p>
        </div>
        <div class="bento-card">
            <div class="bento-icon">🧠</div>
            <h3>AI Digesting</h3>
            <p>Instantly extract Facts, Issues, and Rulings using a multi-provider AI swarm tuned for Philippine legal vernacular.</p>
        </div>
        <div class="bento-card">
            <div class="bento-icon">✅</div>
            <h3>Review Ready</h3>
            <p>Export formatted digests ready for bar review recitations, case files, or legal briefs in Markdown or PDF.</p>
        </div>
    </div>
    """, unsafe_allow_html=True)

# Session state
if "digest" not in st.session_state:
    st.session_state["digest"] = None

# ════════════════════════════════════════════════════════════════════
# HELPERS
# ════════════════════════════════════════════════════════════════════
STAGE_LABELS = [
    "Fetch / Ingest", "Keyword Scan", "Chunk Assembly",
    "Map — AI Swarm", "Stitch & Verify", "Reduce — AI Swarm",
]

def _stages(done, active, msg):
    pills = "".join(
        f'<span class="stage-pill s-done">✓ {l}</span>' if i < done
        else f'<span class="stage-pill s-active">▶ {l}</span>' if i == active
        else f'<span class="stage-pill s-wait">○ {l}</span>'
        for i, l in enumerate(STAGE_LABELS)
    )
    progress_area.markdown(
        f'<div class="stage-row">{pills}</div>'
        f'<div class="stage-msg">{msg}</div>',
        unsafe_allow_html=True,
    )

def _build_swarm():
    import config
    from swarm.provider import ApeKeyProvider
    cfgs = [
        ("Groq",       "https://api.groq.com/openai/v1/chat/completions",                         config.GROQ_API_KEY,       "openai/gpt-oss-120b",            30),
        ("Gemini",     "https://generativelanguage.googleapis.com/v1beta/openai/chat/completions", config.GEMINI_API_KEY,     "gemini-3.5-flash",               15),
        ("Mistral",    "https://api.mistral.ai/v1/chat/completions",                               config.MISTRAL_API_KEY,    "mistral-medium-3-5",             20),
        ("NVIDIA",     "https://integrate.api.nvidia.com/v1/chat/completions",                     config.NVIDIA_API_KEY,     "meta/llama-3.3-70b-instruct",    20),
        ("OpenRouter", "https://openrouter.ai/api/v1/chat/completions",                            config.OPENROUTER_API_KEY, "openai/gpt-oss-20b:free",        20),
    ]
    providers = [
        ApeKeyProvider(name="ApeKey", api_url="https://apekey.ai/v1/chat/completions",
                       api_key=config.APEKEY_API_KEY, model="auto",
                       max_tokens=config.MAP_MAX_TOKENS, rpm_limit=30,
                       timeout=5, prefer="speed"),
    ]
    providers += [
        Provider(name=n, api_url=u, api_key=k, model=m,
                 max_tokens=config.MAP_MAX_TOKENS, rpm_limit=r)
        for n, u, k, m, r in cfgs
    ]
    return ProviderSwarm(providers)

def run_pipeline(source, is_pdf, subject):
    os.makedirs(INTERMEDIATE_DIR, exist_ok=True)
    os.makedirs(os.path.dirname(OUTPUT_DIGEST_PATH), exist_ok=True)
    swarm = _build_swarm()

    _stages(0, 0, "Fetching case text…")
    try:
        lines, gc = ingest_pdf(source) if is_pdf else fetch_case_from_url(source)
        gc["bar_subject"] = subject
    except Exception as e:
        st.error(f"**Stage 0 — Fetch failed:** {e}"); return None

    _stages(1, 1, f"Scanning {len(lines):,} lines…")
    try:
        auto = build_automaton(KEYWORDS)
        hits = deduplicate_hits(scan_lines(lines, auto), min_gap=DEFAULT_MIN_GAP)
        for h in hits:
            if not h["local_verbatim"]:
                h["local_verbatim"] = pre_extract_verbatim(lines, h, window=VERBATIM_WINDOW)
    except Exception as e:
        st.error(f"**Stage 2 — Scan failed:** {e}"); return None

    _stages(2, 2, f"Assembling {len(hits)} chunks…")
    try:
        payloads = build_all_payloads(lines, hits, gc, window=DEFAULT_WINDOW)
    except Exception as e:
        st.error(f"**Stage 3 — Chunk failed:** {e}"); return None

    _stages(3, 3, f"Dispatching {len(payloads)} chunks to swarm…")
    try:
        packets = process_all_chunks(payloads, swarm)
    except Exception as e:
        st.error(f"**Stage 4 — Map failed:** {e}"); return None

    _stages(4, 4, "Stitching and verifying verbatim…")
    try:
        stream = build_compiled_stream(packets, lines, gc)
    except Exception as e:
        st.error(f"**Stage 5 — Stitch failed:** {e}"); return None

    _stages(5, 5, "Running final Reduce pass…")
    try:
        # Bump token limit for the long reduce output (same as run.py)
        for provider in swarm._providers:
            provider.max_tokens = REDUCE_MAX_TOKENS
        digest = run_reduce(stream, gc, swarm)
    except Exception as e:
        st.error(f"**Stage 6 — Reduce failed:** {e}"); return None

    try:
        with open(OUTPUT_DIGEST_PATH, "w", encoding="utf-8") as f:
            f.write(digest)
    except OSError:
        pass

    _stages(6, -1, "Analysis complete.")
    return digest

# ════════════════════════════════════════════════════════════════════
# GENERATE HANDLER
# ════════════════════════════════════════════════════════════════════
if generate_clicked:
    if input_type == "URL":
        if not url_input.strip():
            st.warning("Please enter a case URL."); st.stop()
        source, is_pdf, tmp_pdf = url_input.strip(), False, None
    else:
        if pdf_file is None:
            st.warning("Please upload a PDF."); st.stop()
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
        tmp.write(pdf_file.read()); tmp.flush(); tmp.close()
        source, is_pdf, tmp_pdf = tmp.name, True, tmp.name

    with st.spinner("Running pipeline…"):
        digest = run_pipeline(source, is_pdf, bar_subject)

    if is_pdf and tmp_pdf:
        try: os.unlink(tmp_pdf)
        except OSError: pass

    if digest:
        st.session_state["digest"] = digest
        st.success("✓ Digest generated successfully.")

# ════════════════════════════════════════════════════════════════════
# HELPERS: DOCX conversion
# ════════════════════════════════════════════════════════════════════

def _md_to_docx_bytes(md_text: str) -> bytes:
    """Convert markdown digest text to a .docx file and return as bytes."""
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)

    for line in md_text.split('\n'):
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph('')
            continue

        # Heading: ## Section
        if stripped.startswith('## '):
            p = doc.add_paragraph()
            run = p.add_run(stripped[3:])
            run.bold = True
            run.font.size = Pt(14)
            run.font.color.rgb = None  # default black
            continue

        # Bold headings like **Case:**
        if stripped.startswith('**') and stripped.endswith('**'):
            p = doc.add_paragraph()
            run = p.add_run(stripped.strip('*'))
            run.bold = True
            continue

        # Blockquote: >
        if stripped.startswith('> '):
            p = doc.add_paragraph(stripped[2:])
            p.style = doc.styles['Normal']
            p.paragraph_format.left_indent = Inches(0.5)
            for run in p.runs:
                run.italic = True
            continue

        # Regular paragraph — inline bold/italic
        p = doc.add_paragraph()
        _add_inline_text(p, stripped)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _add_inline_text(paragraph, text: str):
    """Add run segments to a paragraph for inline bold (**) and italic (*) formatting."""
    import re as _re
    # Split on **...** or *...* patterns
    parts = _re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


# ════════════════════════════════════════════════════════════════════
# DIGEST DISPLAY
# ════════════════════════════════════════════════════════════════════
if st.session_state.get("digest"):
    digest_text = st.session_state["digest"]
    st.markdown("---")
    st.markdown("""
    <p style='font-family:"Source Serif 4",serif;font-size:20px;font-weight:700;
              color:#001B3D;margin-bottom:10px;'>Master Digest</p>
    """, unsafe_allow_html=True)

    # Convert markdown → HTML so it renders inside the digest-output div
    digest_html = mistune.html(digest_text)
    output_area.markdown(
        f'<div class="digest-output">{digest_html}</div>',
        unsafe_allow_html=True,
    )

    docx_bytes = _md_to_docx_bytes(digest_text)
    st.download_button(
        "⬇  Export to DOCX",
        data=docx_bytes,
        file_name="master_digest.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

# Footer
st.markdown("""
<div class="lex-footer">
    <span><strong style='color:#D4AF37;font-family:"Hanken Grotesk",sans-serif;'>
        Lex Modernus</strong> &nbsp;·&nbsp; © 2024 Republic of the Philippines</span>
    <span>Privacy Policy &nbsp;·&nbsp; Terms of Service &nbsp;·&nbsp; Official Gazette</span>
</div>
""", unsafe_allow_html=True)
